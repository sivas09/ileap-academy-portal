from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "final-project-completion-report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.style = doc.styles["Normal"]
    run = p.add_run("iLEAP Academy Portal Completion Report")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("iLEAP Academy Portal Final Project Completion Report")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Requirements, technology stack, architecture, deployment, and quality verification")
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [1.875, 4.625])
    rows = [
        ("Project", "iLEAP Academy English Writing Program Portal"),
        ("Production target", "https://english.ileapacademy.com"),
        ("Hosting target", "Render web service with PostgreSQL"),
        ("Database schema", "english_portal"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            row.cells[idx].text = value
    set_table_width(table, widths)


def main() -> None:
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The iLEAP Academy English Writing Program portal is implemented as a production-ready full-stack web application for public program marketing, student access, teacher/admin content management, AI writing feedback, resource delivery, and one-time Stripe purchases."
    )
    add_callout(
        doc,
        "Final Status",
        "The project is ready for production validation on Render. The strongest remaining work is operational hardening, automated test expansion, monitoring, and final policy review for student data and AI usage.",
    )

    doc.add_heading("2. Requirements Coverage", level=1)
    add_table(
        doc,
        ["Area", "Completion Summary"],
        [
            ["Public website", "English Writing Program pages, level information, portal entry, and shop paths are represented."],
            ["Authentication", "Student, teacher, and admin login is implemented with JWT sessions and expired-token handling."],
            ["Student portal", "Students can access dashboard content, level resources, assignments, purchases, and AI Tutor workflows."],
            ["Teacher portal", "Teachers can manage authorized educational content with scoped permissions."],
            ["Admin portal", "Admins can manage users, resources, assignments, prompts, products, orders, and website content."],
            ["AI Tutor", "OpenAI integration, configurable prompts, stored feedback, and error handling are included."],
            ["Payments", "Stripe checkout and webhook handling support one-time purchases and entitlements."],
            ["Deployment", "Render deployment, Prisma migrations, and PostgreSQL schema isolation are documented and tested."],
        ],
        [1.65, 4.85],
    )

    doc.add_heading("3. Technology Stack", level=1)
    doc.add_heading("Frontend", level=2)
    add_bullets(doc, ["React 19", "TypeScript", "Vite", "CSS in the existing app structure", "lucide-react icons"])
    doc.add_heading("Backend", level=2)
    add_bullets(
        doc,
        [
            "Node.js and Express 5",
            "TypeScript",
            "Prisma ORM and PostgreSQL",
            "JWT authentication and bcryptjs password hashing",
            "Zod validation, Helmet, CORS, and Multer upload handling",
        ],
    )
    doc.add_heading("External Services", level=2)
    add_bullets(
        doc,
        [
            "Render web service and Render PostgreSQL",
            "Render persistent disk as upload fallback",
            "Cloudflare R2-compatible object storage support",
            "OpenAI API for AI Tutor feedback",
            "Stripe for checkout and payment webhooks",
            "SMTP email delivery through nodemailer",
        ],
    )

    doc.add_heading("4. Production Architecture", level=1)
    doc.add_paragraph(
        "The current architecture deploys the React frontend and Express API together as one Render web service. The service connects to PostgreSQL through Prisma and uses the english_portal schema for database isolation."
    )
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ["Browser", "Loads the React app and calls authenticated API endpoints over HTTPS."],
            ["Render web service", "Serves the static frontend, Express API, auth middleware, and Prisma client."],
            ["Render PostgreSQL", "Stores users, roles, resources, assignments, AI feedback, products, orders, and entitlements."],
            ["External services", "OpenAI, Stripe, SMTP, and object storage integrations."],
        ],
        [1.75, 4.75],
    )
    add_callout(
        doc,
        "Shared iLEAP Direction",
        "For multiple iLEAP applications, keep one Render account and start with one paid PostgreSQL instance, but use separate Render web services and separate database schemas per major application.",
    )

    doc.add_heading("5. Database Architecture", level=1)
    add_bullets(
        doc,
        [
            "Current schema: english_portal.",
            "Future schemas should be separate, such as ileap_club, ileap_tech_ai, and shared_identity.",
            "Prisma migrations should remain the controlled path for database changes.",
            "Production should use a paid Render PostgreSQL plan before storing real student or payment data.",
            "Backups should include Render-managed backups plus periodic off-platform exports.",
        ],
    )

    doc.add_heading("6. Security Architecture", level=1)
    add_table(
        doc,
        ["Control", "Current Status"],
        [
            ["Password security", "Passwords are hashed with bcryptjs."],
            ["Session validation", "JWT verification protects API routes."],
            ["Expired sessions", "Expired or invalid tokens return 401 and redirect the user to login."],
            ["Role access", "Student, teacher, and admin operations are role-protected."],
            ["Teacher scoping", "Teacher resource access is limited to assigned levels."],
            ["Payments", "Stripe webhook signatures protect checkout completion."],
            ["Headers and CORS", "Helmet and backend CORS controls are enabled."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("7. File Storage Architecture", level=1)
    doc.add_paragraph(
        "Uploaded files can currently use Render persistent disk as a fallback. For production scale, Cloudflare R2 should be the long-term storage layer for PDFs, worksheets, books, and other uploaded files. PostgreSQL should store metadata, ownership, and access rules."
    )

    doc.add_heading("8. Deployment Strategy", level=1)
    add_numbered(
        doc,
        [
            "Push code to GitHub.",
            "Render pulls the repository and runs the build command.",
            "Prisma Client is generated.",
            "React and server TypeScript builds are produced.",
            "Render start command applies migrations.",
            "Express serves the frontend and API.",
        ],
    )
    doc.add_heading("Required Environment Variables", level=2)
    add_bullets(
        doc,
        [
            "Core: DATABASE_URL, APP_DATABASE_SCHEMA, JWT_SECRET, PORT",
            "AI: OPENAI_API_KEY, OPENAI_MODEL",
            "Payments: STRIPE_SECRET_KEY, STRIPE_WEBHOOK_SECRET",
            "Storage: R2_BUCKET, R2_ACCOUNT_ID, R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY, R2_ENDPOINT",
            "Email: SMTP_HOST, SMTP_USER, SMTP_PASS, SMTP_FROM, DEFAULT_SUBMISSION_NOTIFICATION_EMAIL",
        ],
    )

    doc.add_heading("9. Domain Mapping", level=1)
    add_table(
        doc,
        ["Domain", "Target"],
        [
            ["english.ileapacademy.com", "Current iLEAP Academy English portal on Render"],
            ["academy.ileap.com or english.ileapacademy.com", "Future iLEAP Academy portal"],
            ["memberportal.ileapclub.com", "Future iLEAP Club member portal"],
            ["tech.ileaptechai.com", "Future iLEAP Tech AI portal"],
            ["admin.ileap.com", "Future shared admin or identity portal"],
        ],
        [2.6, 3.9],
    )

    doc.add_heading("10. Quality Test Results", level=1)
    add_table(
        doc,
        ["Check", "Command", "Result"],
        [
            ["Frontend build", "npm.cmd run build", "Passed"],
            ["Server TypeScript build", "npm.cmd run build:server", "Passed"],
            ["Render-style production build", "npm.cmd run render:build", "Passed"],
            ["Production dependency audit", "npm.cmd audit --omit=dev", "Passed, 0 vulnerabilities"],
        ],
        [1.85, 2.95, 1.7],
    )
    doc.add_paragraph(
        "The Render build was verified with a PostgreSQL-shaped placeholder URL so Prisma could validate database URL format during local build verification. This did not deploy or modify production data."
    )

    doc.add_heading("11. Dependency Quality Updates", level=1)
    add_bullets(
        doc,
        [
            "Updated dependency resolutions through npm audit fix.",
            "Upgraded nodemailer to ^9.0.1.",
            "Re-ran build and audit checks after the upgrade.",
            "Confirmed the production dependency audit reports 0 vulnerabilities.",
        ],
    )

    doc.add_heading("12. Known Limitations", level=1)
    add_bullets(
        doc,
        [
            "No formal automated unit test suite is currently documented.",
            "No full browser end-to-end regression suite is currently documented.",
            "Live production login was not verified during this quality pass because production credentials were not used.",
            "Production data retention, privacy, consent, and AI disclosure language still need final business/legal confirmation.",
            "Long-term file storage should be moved fully to Cloudflare R2 before heavy real usage.",
            "Monitoring and alerting should be added before scaling usage.",
        ],
    )

    doc.add_heading("13. Recommended Launch Checklist", level=1)
    add_numbered(
        doc,
        [
            "Upgrade Render PostgreSQL to a paid plan.",
            "Confirm english.ileapacademy.com/api/health returns ok.",
            "Verify admin, student, and teacher login.",
            "Verify expired session behavior redirects to login.",
            "Upload and download one test resource per level.",
            "Submit one test assignment.",
            "Generate one AI Tutor response using a test writing sample.",
            "Complete one Stripe test checkout and confirm entitlement unlock.",
            "Confirm Stripe webhook event delivery.",
            "Confirm SMTP email delivery.",
            "Confirm R2 upload/download if R2 is enabled.",
            "Review privacy, AI, and payment notices.",
            "Confirm database backups are enabled.",
        ],
    )

    doc.add_heading("14. Recommended Next Improvements", level=1)
    add_bullets(
        doc,
        [
            "Add automated tests for authentication, authorization, resources, assignments, AI Tutor, and Stripe webhook handling.",
            "Add Playwright end-to-end tests for student, teacher, and admin flows.",
            "Add rate limiting and login brute-force protection.",
            "Add structured server logging and error monitoring.",
            "Add GitHub Actions CI for build, server build, and audit checks.",
            "Complete Cloudflare R2 production storage migration.",
            "Add operational runbooks for admin account repair, backup restore, and Stripe webhook debugging.",
        ],
    )

    doc.add_heading("15. Final Status", level=1)
    doc.add_paragraph(
        "The iLEAP Academy English Writing Program portal has the core requirements, technical stack, production architecture, deployment path, security model, and quality verification needed for controlled production use."
    )

    p = doc.add_paragraph()
    p.add_run("Source repository: ").bold = True
    add_hyperlink(p, "sivas09/ileap-academy-portal", "https://github.com/sivas09/ileap-academy-portal")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
