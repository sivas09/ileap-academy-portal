from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE = DOCS / "english-writing-portal-requirements.md"
OUTPUT = DOCS / "iLEAP_English_Writing_Portal_Requirements.docx"
LOGO = ROOT / "Logo_large.jpg"

BRAND = {
    "background": "FFEEFB",
    "red": "C62828",
    "blue": "00296B",
    "green": "008000",
    "orange": "E6690A",
    "light_blue": "0089CA",
    "ink": "1F2937",
    "muted": "5B6472",
    "soft_blue": "EAF6FC",
    "soft_red": "FDEAEA",
    "soft_green": "EAF7EA",
}


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE8", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_page_background(document, fill):
    settings = document.settings.element
    background = OxmlElement("w:background")
    background.set(qn("w:color"), fill)
    settings.insert(0, background)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BRAND["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 24, BRAND["blue"], 0, 8),
        ("Subtitle", 11, BRAND["muted"], 0, 12),
        ("Heading 1", 16, BRAND["blue"], 14, 6),
        ("Heading 2", 13, BRAND["light_blue"], 10, 5),
        ("Heading 3", 11.5, BRAND["orange"], 8, 3),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")


def add_header_footer(document):
    section = document.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO.exists():
        run = hp.add_run()
        run.add_picture(str(LOGO), width=Inches(0.8))

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("iLEAP Academy | English Writing Program Portal Requirements")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(BRAND["muted"])


def add_cover(document):
    if LOGO.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.5))

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("English Writing Program Portal Requirements")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Functional and non-functional requirements for the redesigned iLEAP Academy website and student portal")

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    for row in meta.rows:
        row.cells[0].width = Inches(1.45)
        row.cells[1].width = Inches(4.65)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, bottom=90)
    set_table_borders(meta, color="E2E8F0")
    values = [
        ("Version", "1.0"),
        ("Date", "May 19, 2026"),
        ("Scope", "MVP requirements for English Writing Program only"),
        ("Users", "Students, teachers, and admins"),
    ]
    for row, (k, v) in zip(meta.rows, values):
        shade_cell(row.cells[0], BRAND["soft_blue"])
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)

    document.add_paragraph()
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Brand basis: iLEAP Academy logo and 2026 branding guide in the project workspace.")
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(BRAND["muted"])

    document.add_page_break()


def add_callout(document, title, text, fill):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    set_table_borders(table, color="D9DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BRAND["blue"])
    p.add_run("\n" + text)
    document.add_paragraph()


def add_level_table(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.65), Inches(1.8), Inches(3.05)]
    headers = ["Level", "Program Name", "Primary Writing Focus"]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].width = widths[idx]
        shade_cell(table.rows[0].cells[idx], BRAND["blue"])
        p = table.rows[0].cells[idx].paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    rows = [
        ("Grade 2/3", "Story Builder", "Sentences, paragraph foundations, creative story writing, early EQAO-style preparation."),
        ("Grade 4/5/6", "Paragraph Builder", "Paragraph structure, topic sentences, five-paragraph essay foundations, EQAO preparation."),
        ("Grade 7/8/9", "Essay Mastery", "Persuasive and analytical essays, thesis development, research writing, OSSLT-style timed writing."),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].width = widths[idx]
            cells[idx].text = text
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cells[0], BRAND["soft_red"] if row_data[0] == "Grade 2/3" else BRAND["soft_green"] if row_data[0] == "Grade 4/5/6" else BRAND["soft_blue"])
    set_table_borders(table)
    document.add_paragraph()


def add_requirements_table(document, heading, prefix, items):
    document.add_heading(heading, level=2)
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.8), Inches(4.65), Inches(1.05)]
    for idx, text in enumerate(["ID", "Requirement", "Priority"]):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        shade_cell(cell, BRAND["blue"])
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for idx, text in enumerate(items, start=1):
        cells = table.add_row().cells
        values = [f"{prefix}-{idx:02d}", text, "MVP"]
        for col, value in enumerate(values):
            cells[col].width = widths[col]
            cells[col].text = value
            set_cell_margins(cells[col], top=90, bottom=90)
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if col in (0, 2):
                cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cells[0], "F8FAFC")
    set_table_borders(table)


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.add_run(text)


def add_numbered(document, text):
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.28)
    p.add_run(text)


def add_markdown_body(document):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    skip_until = "## 1. Purpose"
    started = False
    current_list = None
    for raw in lines:
        line = raw.strip()
        if line == skip_until:
            started = True
        if not started:
            continue
        if not line:
            current_list = None
            continue

        if line == "## 5. Program Levels":
            document.add_heading("5. Program Levels", level=1)
            add_level_table(document)
            current_list = None
            continue

        if line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            current_list = None
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            current_list = None
            continue
        if line.startswith("- "):
            add_bullet(document, line[2:])
            current_list = "bullet"
            continue

        req_match = re.match(r"^(FR|NFR)-\d+:\s+(.*)$", line)
        if req_match:
            p = document.add_paragraph()
            r = p.add_run(req_match.group(1) + ": ")
            r.bold = True
            p.add_run(req_match.group(2))
            continue

        p = document.add_paragraph()
        p.add_run(line)


def build_document():
    document = Document()
    style_document(document)
    set_page_background(document, BRAND["background"])
    add_header_footer(document)
    add_cover(document)

    add_callout(
        document,
        "MVP Decision",
        "This document focuses only on the English Writing Program portal: public marketing, student dashboard, assignments, AI tutor feedback, resource purchases, Stripe checkout, and admin management.",
        BRAND["soft_blue"],
    )
    add_markdown_body(document)

    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
